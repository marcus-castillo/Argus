"""Extract plain text from uploaded PDF / DOCX / TXT files."""
from __future__ import annotations

import io

from app.core.errors import UnsupportedFileTypeError

PDF_TYPES = {"application/pdf"}
DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
TEXT_TYPES = {"text/plain", "text/markdown", "application/octet-stream"}

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class ExtractedDocument:
    __slots__ = ("text", "page_count")

    def __init__(self, text: str, page_count: int | None) -> None:
        self.text = text
        self.page_count = page_count


def extract_text(filename: str, content_type: str, data: bytes) -> ExtractedDocument:
    """Dispatch on content type / extension and return extracted text."""
    lower = filename.lower()
    if content_type in PDF_TYPES or lower.endswith(".pdf"):
        return _extract_pdf(data)
    if content_type in DOCX_TYPES or lower.endswith(".docx"):
        return _extract_docx(data)
    if content_type in TEXT_TYPES or lower.endswith((".txt", ".md")):
        return ExtractedDocument(data.decode("utf-8", errors="replace"), None)
    raise UnsupportedFileTypeError(
        f"Unsupported file type: {content_type or filename}. "
        "Upload a PDF, DOCX, or plain-text file."
    )


def _extract_pdf(data: bytes) -> ExtractedDocument:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = [page.extract_text() or "" for page in reader.pages]
    return ExtractedDocument("\n".join(parts), page_count=len(reader.pages))


def _extract_docx(data: bytes) -> ExtractedDocument:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    parts: list[str] = [para.text for para in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return ExtractedDocument("\n".join(parts), page_count=None)
